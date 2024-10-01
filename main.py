import os, sys, json, re, subprocess, chardet
from helper import *
from translationHelper import *
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from langdetect.lang_detect_exception import LangDetectException

SUCCESS_INDICATOR = "    ✔️ "
ERROR_INDICATOR = "    ❌ "
def set_default_font(doc, font_name="Noto Sans"):
    """Set the default font for the entire document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    # Handle the case where the font needs to be set specifically for East Asian languages
    font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_docx_from_structure(output_path, structure, text_files):
    try:
        # Generate timestamp and output paths
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_docx_path = os.path.join(output_path, f"{timestamp}.docx")

        # Create a new DOCX document
        doc = Document()

        # Set the default font to Noto Sans
        set_default_font(doc)

        # Helper function to add paragraphs
        def add_paragraph(doc, text, alignment, font_size, bold=False):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            paragraph.alignment = alignment

        # Accumulate all translated content to add at the end
        translation_section = []

        # Process each item in the structure list
        for item in structure:
            if item['type'] == 'heading':
                add_paragraph(doc, item['text'], WD_ALIGN_PARAGRAPH.CENTER, item['fontSize'], bold=True)
            elif item['type'] == 'normal':
                add_paragraph(doc, item['text'], WD_ALIGN_PARAGRAPH.JUSTIFY, item['fontSize'])
            elif item['type'] == 'pagebreak':
                doc.add_page_break()
            elif item['type'] == 'combinedText':
                # Insert the content from each text file
                for text_file in text_files:

                    with open(text_file, 'rb') as f:
                        raw_data = f.read()
                        result = chardet.detect(raw_data)
                        encoding = result['encoding']

                    with open(text_file, 'r', encoding=encoding) as f:
                        content = f.readlines()
                    
                        # Remove the first line if it matches the unwanted text
                        if content[0].strip() == "OCR/HTR":
                            content = content[1:]  # Remove the first line
                        
                        # Proceed as before with the updated content list
                        page_heading = content[0].strip()  # First line (page heading)
                        # page_heading = page_heading.replace(" ", "-")
                        text_content = ''.join(content[1:])  # Remaining text

                        # Remove leading newline characters if they exist
                        if text_content.startswith('\n'):
                            text_content = text_content.lstrip('\n')

                        # Check for non-English content and translate it immediately
                        contains_non_english, translated_text = check_and_translate_non_english_content(text_content)

                        # Left-align the first line (page heading) and make it bold
                        add_paragraph(doc, page_heading, WD_ALIGN_PARAGRAPH.LEFT, item['fontSize'] + 2, bold=True)
                        # Justify the rest of the text
                        add_paragraph(doc, text_content, WD_ALIGN_PARAGRAPH.JUSTIFY, item['fontSize'])

                        # If non-English content is detected, accumulate the translation
                        if contains_non_english:
                            translation_section.append({
                                'heading': page_heading,
                                'translated_text': translated_text,
                                'font_size': item['fontSize']
                            })

                        # Add a newline after each file's content
                        doc.add_paragraph()

        # Add the translation section at the end of the document
        if translation_section:
            add_paragraph(doc, "Translation", WD_ALIGN_PARAGRAPH.CENTER, item['fontSize'] + 4, bold=True)
            for section in translation_section:
                add_paragraph(doc, section['heading'], WD_ALIGN_PARAGRAPH.LEFT, section['font_size'] + 2, bold=True)
                add_paragraph(doc, section['translated_text'], WD_ALIGN_PARAGRAPH.JUSTIFY, section['font_size'])
                add_paragraph(doc, "", WD_ALIGN_PARAGRAPH.JUSTIFY, section['font_size'])

        # Save the DOCX file
        doc.save(output_docx_path)
        return output_docx_path
    except Exception as e:
        print(f"{ERROR_INDICATOR}Error creating DOCX: {e}")
        exit(1)

def run_script(args, config):
    # Get the root folder directory from the configuration
    root_folder = config['root_folder_directory']
    output_path = config['output_folder_directory']
    
    print(f"Root folder directory: {root_folder}")
    print(f"Output folder directory: {output_path}\n\n")
    
    os.makedirs(output_path, exist_ok=True)

    # Iterate through all subfolders in the root directory
    for subfolder in os.listdir(root_folder):
        subfolder_path = os.path.join(root_folder, subfolder)
        print(f"Checking subfolder: {subfolder_path}")

        # Check if it's a directory
        if os.path.isdir(subfolder_path):
            # Initialize the path for the OCR results folder
            ocr_results_folder = None

            # Look for any folder that ends with '_ocr_results'
            for folder in os.listdir(subfolder_path):
                folder_path = os.path.join(subfolder_path, folder)
                if os.path.isdir(folder_path) and folder.endswith('_ocr_results'):
                    ocr_results_folder = folder_path
                    print(f"OCR results folder found: {ocr_results_folder}")
                    break  # Stop searching once we find the first match

            # Check if the OCR results folder was found
            if ocr_results_folder:
                # Get the base name of the OCR results folder (excluding "_ocr_results")
                sibling_folder_name = ocr_results_folder.replace('_ocr_results', '')  # Remove "_ocr_results" from the folder name
                sibling_folder = os.path.join(subfolder_path, sibling_folder_name)  # Construct sibling folder path
                
                print(f"Sibling folder determined: {sibling_folder}")

                # Collect all .txt files from the OCR results folder
                text_files = [os.path.join(ocr_results_folder, f) for f in os.listdir(ocr_results_folder) if f.endswith('.txt')]
                
                print(f"Found {len(text_files)} .txt files in {ocr_results_folder}.")
                
                # If no text files found, skip to the next subfolder
                if not text_files:
                    print(f"No .txt files found in {ocr_results_folder}. Skipping...\n")
                    continue

                # Create DOCX and convert to PDF
                print(f"Creating DOCX from {len(text_files)} text files...")
                docx_file = create_docx_from_structure(output_path, config['doc_structure'], text_files)
                print(f"DOCX created at: {docx_file}")

                print("Converting DOCX to PDF...")
                pdf_file = convert_docx_to_pdf(docx_file)
                print(f"PDF created at: {pdf_file}\n")

                # Ensure the sibling folder exists
                os.makedirs(sibling_folder, exist_ok=True)
                print(f"Ensured sibling folder exists: {sibling_folder}")

                # Merge PDFs into the sibling folder
                try:
                    print(f"Merging PDF into sibling folder: {sibling_folder}...")
                    subprocess.run([sys.executable, 'pdf_merger.py', pdf_file, sibling_folder], check=True)
                    print(f"PDF successfully merged into: {sibling_folder}\n")
                except subprocess.CalledProcessError as e:
                    print(f"{ERROR_INDICATOR} Error occurred while merging PDFs: {e}\n")
            else:
                print(f"No OCR results folder found for {subfolder}. Skipping...\n")
        else:
            print(f"{subfolder_path} is not a directory. Skipping...\n")

def main(args=None):
    config_path = 'config.json'
    config = load_config(config_path)

    while True:
        user_choice = display_menu()
        clear_terminal()
        if user_choice == '1':
            run_script(args, config)
        elif user_choice == '2':
            change_directory(config, config_path)
        elif user_choice == '3':
            change_directory(config, config_path, False)
        elif user_choice == '4':
            break
        else:
            print("Invalid choice. Please enter 1, 2, or 3.")
        
        wait_for_keypress()

if __name__ == "__main__":
    main(sys.argv)
